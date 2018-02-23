#! python3
# invites.py - A program that generates custom invitations to guests

import docx

file = open('guests.txt')
guestNames = file.readlines()

doc = docx.Document('C:\\Users\\Rodrigo\\style.docx')
os.chdir('C:\\Users\\Rodrigo')
for name in guestNames:
	doc.paragraphs[0].text = 'It would be a pleasure to have the company of'
	doc.paragraphs[1].text = guestNames[name]
	doc.paragraphs[2].text = 'at 2020 Memory Lane on the Evening of'
	doc.paragraphs[3].text = 'April 1st'
	doc.paragraphs[4].text = "9 o'clock"
	if name != guestNames[-1]:
		doc.paragraphs[4].runs[0].add_break(docx.text.WD_BREAK.PAGE)